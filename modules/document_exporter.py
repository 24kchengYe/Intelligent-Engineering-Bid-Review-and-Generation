"""
文档导出模块
支持将Markdown格式的报告和技术标转换为Word文档
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from typing import Dict, List


class MarkdownToWordConverter:
    """Markdown转Word转换器（保持排版）"""

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """设置Word样式"""
        # 设置中文字体
        self.doc.styles['Normal'].font.name = '宋体'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        self.doc.styles['Normal'].font.size = Pt(12)

        # 设置标题样式
        for i in range(1, 4):
            heading_style = self.doc.styles[f'Heading {i}']
            heading_style.font.name = '黑体'
            heading_style.font.bold = True
            heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

            # 标题大小
            if i == 1:
                heading_style.font.size = Pt(18)
            elif i == 2:
                heading_style.font.size = Pt(16)
            else:
                heading_style.font.size = Pt(14)

    def convert(self, markdown_text: str) -> Document:
        """
        将Markdown文本转换为Word文档

        Args:
            markdown_text: Markdown格式的文本

        Returns:
            docx.Document对象
        """
        lines = markdown_text.split('\n')
        in_code_block = False
        in_table = False
        table_lines = []

        for line in lines:
            # 代码块
            if line.startswith('```'):
                in_code_block = not in_code_block
                continue

            if in_code_block:
                # 代码块内容（等宽字体）
                p = self.doc.add_paragraph(line)
                p.style = 'Normal'
                p.runs[0].font.name = 'Courier New'
                p.runs[0].font.size = Pt(10)
                continue

            # 标题（支持 # ## ### ）
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                title_text = line.lstrip('#').strip()

                if level <= 3:
                    self.doc.add_heading(title_text, level=level)
                else:
                    # 超过3级的标题用加粗段落
                    p = self.doc.add_paragraph(title_text)
                    p.runs[0].bold = True
                    p.runs[0].font.size = Pt(13)
                continue

            # 分割线
            if line.strip() in ['---', '***', '___']:
                self.doc.add_paragraph('_' * 80)
                continue

            # 表格检测
            if '|' in line and line.strip().startswith('|'):
                if not in_table:
                    in_table = True
                    table_lines = []
                table_lines.append(line)
                continue
            elif in_table and line.strip() == '':
                # 表格结束
                self._add_table(table_lines)
                in_table = False
                table_lines = []
                continue

            # 列表（无序）
            if re.match(r'^\s*[-*+]\s+', line):
                text = re.sub(r'^\s*[-*+]\s+', '', line)
                self._add_paragraph_with_formatting(text, style='List Bullet')
                continue

            # 列表（有序）
            if re.match(r'^\s*\d+\.\s+', line):
                text = re.sub(r'^\s*\d+\.\s+', '', line)
                self._add_paragraph_with_formatting(text, style='List Number')
                continue

            # 普通段落
            if line.strip():
                self._add_paragraph_with_formatting(line.strip())

        # 处理最后的表格
        if in_table and table_lines:
            self._add_table(table_lines)

        return self.doc

    def _add_paragraph_with_formatting(self, text: str, style='Normal'):
        """添加段落并处理内联格式（加粗、斜体）"""
        p = self.doc.add_paragraph(style=style)

        # 处理加粗 **text** 和 __text__
        parts = re.split(r'(\*\*.*?\*\*|__.*?__)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # 加粗
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('__') and part.endswith('__'):
                # 加粗
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                # 斜体
                run = p.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('_') and part.endswith('_') and not part.startswith('__'):
                # 斜体
                run = p.add_run(part[1:-1])
                run.italic = True
            else:
                p.add_run(part)

    def _add_table(self, table_lines: List[str]):
        """添加表格"""
        if not table_lines:
            return

        # 解析表格
        rows = []
        for line in table_lines:
            # 跳过分隔线 |---|---|
            if re.match(r'\|[\s\-:]+\|', line):
                continue

            # 提取单元格
            cells = [cell.strip() for cell in line.split('|')]
            # 移除开头和结尾的空元素
            cells = [c for c in cells if c]
            if cells:
                rows.append(cells)

        if not rows:
            return

        # 创建Word表格
        num_cols = max(len(row) for row in rows)
        table = self.doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Light Grid Accent 1'

        # 填充内容
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = cell_text

                    # 表头加粗
                    if row_idx == 0:
                        cell.paragraphs[0].runs[0].font.bold = True

        # 表格后添加空行
        self.doc.add_paragraph()

    def save(self, file_path: str):
        """保存Word文档"""
        self.doc.save(file_path)


class DocumentExporter:
    """文档导出器（支持TXT、MD、Word）"""

    @staticmethod
    def export_to_txt(content: str, file_path: str):
        """导出为纯文本"""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)

    @staticmethod
    def export_to_markdown(content: str, file_path: str):
        """导出为Markdown"""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)

    @staticmethod
    def export_to_word(markdown_content: str, file_path: str, title: str = None):
        """
        将Markdown内容导出为Word文档

        Args:
            markdown_content: Markdown格式的内容
            file_path: 保存路径
            title: 文档标题（可选）
        """
        converter = MarkdownToWordConverter()

        # 添加标题
        if title:
            converter.doc.add_heading(title, level=0)
            converter.doc.add_paragraph()

        # 转换内容
        converter.convert(markdown_content)

        # 保存
        converter.save(file_path)

    @staticmethod
    def create_technical_proposal_word(
        outline: Dict,
        generated_sections: Dict[str, str],
        project_name: str = "技术标文档"
    ) -> str:
        """
        创建完整的技术标Word文档

        Args:
            outline: 技术标目录结构
            generated_sections: 已生成的章节内容
            project_name: 项目名称

        Returns:
            Word文档的临时路径
        """
        from datetime import datetime
        import tempfile

        doc = Document()

        # 设置字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        doc.styles['Normal'].font.size = Pt(12)

        # 封面
        title = doc.add_heading(project_name, level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        subtitle_run = subtitle.add_run("技术标文件")
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.name = '黑体'

        doc.add_paragraph()
        date_para = doc.add_paragraph(f"生成日期: {datetime.now().strftime('%Y年%m月%d日')}")
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 分页
        doc.add_page_break()

        # 目录（简化版）
        doc.add_heading("目录", level=1)
        if not outline.get('raw'):
            DocumentExporter._add_outline_to_word(doc, outline)
        doc.add_page_break()

        # 添加章节内容
        for section_title, content in generated_sections.items():
            # 章节标题
            doc.add_heading(section_title, level=1)

            # 转换Markdown内容到Word
            converter = MarkdownToWordConverter()
            converter.doc = doc
            converter.convert(content)

            # 章节后分页
            doc.add_page_break()

        # 保存到临时文件
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()

        return temp_file.name

    @staticmethod
    def _add_outline_to_word(doc, outline_data):
        """将目录结构添加到Word"""
        if 'outline' in outline_data and isinstance(outline_data['outline'], list):
            for item in outline_data['outline']:
                DocumentExporter._add_outline_item(doc, item, level=1)

    @staticmethod
    def _add_outline_item(doc, item, level=1):
        """递归添加目录项"""
        title = item.get('title', '')
        word_count = item.get('word_count', '')

        # 添加目录项
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25 * (level - 1))

        if word_count:
            p.add_run(f"{title} (建议{word_count}字)")
        else:
            p.add_run(title)

        # 递归添加子项
        if 'children' in item and item['children']:
            for child in item['children']:
                DocumentExporter._add_outline_item(doc, child, level + 1)
