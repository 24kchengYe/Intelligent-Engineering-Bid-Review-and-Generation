"""
文档解析模块
支持 PDF、Word、Excel 等格式的文件解析
增强功能：
- 扫描版PDF自动OCR识别
- PDF表格结构化提取
"""

import fitz  # PyMuPDF
from docx import Document
import openpyxl
import pandas as pd
from typing import Dict, Optional, List
import os

# OCR支持（可选）
try:
    from rapidocr_onnxruntime import RapidOCR
    HAS_OCR = True
except ImportError:
    HAS_OCR = False


class DocumentParser:
    """文档解析器（支持OCR和表格结构化）"""

    def __init__(self, enable_ocr=True, extract_tables=True):
        """
        Args:
            enable_ocr: 是否启用OCR识别扫描版PDF
            extract_tables: 是否提取PDF表格结构（可能很慢）
        """
        self.supported_formats = {
            'pdf': self.parse_pdf,
            'docx': self.parse_word,
            'doc': self.parse_word,
            'xlsx': self.parse_excel,
            'xls': self.parse_excel
        }
        self.enable_ocr = enable_ocr and HAS_OCR
        self.extract_tables = extract_tables
        self._ocr_engine = None  # 延迟初始化

    def parse(self, file_path: str, progress_callback=None) -> Dict[str, str]:
        """
        解析文档

        Args:
            file_path: 文件路径
            progress_callback: 进度回调函数（可选）

        Returns:
            解析结果字典，包含 content（文本内容）和 metadata（元数据）
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        # 获取文件扩展名
        file_ext = os.path.splitext(file_path)[1].lower().strip('.')

        if file_ext not in self.supported_formats:
            raise ValueError(f"不支持的文件格式: {file_ext}")

        # 调用对应的解析方法
        parser_func = self.supported_formats[file_ext]

        # PDF支持progress_callback，其他格式不支持
        if file_ext == 'pdf' and progress_callback:
            return parser_func(file_path, progress_callback=progress_callback)
        else:
            return parser_func(file_path)

    def _get_ocr_engine(self):
        """延迟初始化OCR引擎"""
        if self._ocr_engine is None and HAS_OCR:
            self._ocr_engine = RapidOCR()
        return self._ocr_engine

    def _is_scanned_page(self, page) -> bool:
        """判断PDF页面是否为扫描页（无文本层）"""
        text = page.get_text("text")
        return not text or len(text.strip()) < 30

    def _ocr_page(self, page, page_num=0) -> str:
        """使用OCR识别PDF页面（带超时保护）"""
        import signal

        def timeout_handler(signum, frame):
            raise TimeoutError("OCR识别超时")

        try:
            ocr_engine = self._get_ocr_engine()
            if ocr_engine is None:
                return ""

            # 渲染页面为图片（2倍DPI提高识别率）
            mat = fitz.Matrix(2, 2)
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")

            # OCR识别（带超时保护：30秒）
            try:
                # Windows不支持signal.alarm，使用线程超时
                import threading

                result_holder = [None]
                error_holder = [None]

                def ocr_worker():
                    try:
                        res, _ = ocr_engine(img_bytes)
                        result_holder[0] = res
                    except Exception as e:
                        error_holder[0] = e

                thread = threading.Thread(target=ocr_worker)
                thread.daemon = True
                thread.start()
                thread.join(timeout=30)  # 30秒超时

                if thread.is_alive():
                    print(f"⚠️ 第{page_num+1}页OCR识别超时(>30秒)，跳过此页")
                    return f"[第{page_num+1}页OCR识别超时，内容可能缺失]"

                if error_holder[0]:
                    raise error_holder[0]

                result = result_holder[0]

            except Exception as e:
                print(f"⚠️ 第{page_num+1}页OCR识别失败: {e}")
                return f"[第{page_num+1}页OCR识别失败]"

            if result:
                # 提取文本，过滤低置信度（<0.5）
                texts = []
                for item in result:
                    if len(item) >= 3:
                        text, confidence = item[1], item[2]
                        if confidence >= 0.5:
                            texts.append(text)
                    elif len(item) >= 2:
                        texts.append(item[1])
                return "\n".join(texts)
            return ""
        except Exception as e:
            print(f"第{page_num+1}页OCR识别失败: {e}")
            return f"[第{page_num+1}页OCR识别失败]"

    def _extract_tables(self, page) -> List[List[List[str]]]:
        """提取PDF页面中的表格（结构化）"""
        tables = []
        try:
            # PyMuPDF 1.23+ 支持表格检测
            if hasattr(page, 'find_tables'):
                tabs = page.find_tables()
                for tab in tabs:
                    table_data = tab.extract()
                    if table_data:
                        tables.append(table_data)
        except Exception:
            pass
        return tables

    def parse_pdf(self, file_path: str, progress_callback=None) -> Dict[str, str]:
        """解析 PDF 文件（支持扫描版OCR + 表格结构化）

        Args:
            file_path: PDF文件路径
            progress_callback: 进度回调函数 callback(page_num, total_pages, message)
        """
        try:
            doc = fitz.open(file_path)
            content = []
            all_tables = []
            ocr_pages = 0
            ocr_failed_pages = []
            total_pages = len(doc)

            for page_num in range(total_pages):
                page = doc[page_num]

                # 进度回调
                if progress_callback:
                    progress_callback(page_num + 1, total_pages, f"正在处理第 {page_num+1}/{total_pages} 页")

                # 1. 提取文本（优先文本层，降级OCR）
                text = page.get_text("text")
                is_scanned = self._is_scanned_page(page)

                if is_scanned and self.enable_ocr:
                    # 扫描页：使用OCR（带超时保护）
                    print(f"[OCR] 正在识别第 {page_num+1}/{total_pages} 页...")
                    if progress_callback:
                        progress_callback(page_num + 1, total_pages, f"OCR识别第 {page_num+1}/{total_pages} 页...")

                    ocr_text = self._ocr_page(page, page_num)
                    if ocr_text and not ocr_text.startswith("[第"):
                        text = ocr_text
                        ocr_pages += 1
                    elif ocr_text.startswith("[第"):
                        # OCR失败或超时
                        ocr_failed_pages.append(page_num + 1)
                        text = ocr_text  # 保留失败标记

                # 2. 提取表格（结构化）- 可选，因为很慢
                if self.extract_tables:
                    tables = self._extract_tables(page)
                    if tables:
                        all_tables.extend([{
                            'page': page_num + 1,
                            'data': table
                        } for table in tables])

                        # 将表格转为文本添加到内容中
                        for idx, table in enumerate(tables):
                            table_text = self._table_to_text(table)
                            text += f"\n\n[表格 {idx+1}]\n{table_text}"

                if text.strip():
                    content.append(f"--- 第 {page_num + 1} 页 ---\n{text}")

            doc.close()

            metadata = {
                'type': 'PDF',
                'pages': total_pages,
                'file_name': os.path.basename(file_path)
            }

            if ocr_pages > 0:
                metadata['ocr_pages'] = ocr_pages
                metadata['ocr_enabled'] = True

            if ocr_failed_pages:
                metadata['ocr_failed_pages'] = ocr_failed_pages
                metadata['ocr_failed_count'] = len(ocr_failed_pages)

            if all_tables:
                metadata['tables_count'] = len(all_tables)
                metadata['tables'] = all_tables

            return {
                'content': '\n\n'.join(content),
                'metadata': metadata
            }
        except Exception as e:
            return {
                'content': '',
                'metadata': {'type': 'PDF', 'error': str(e)}
            }

    def _table_to_text(self, table_data: List[List[str]]) -> str:
        """将表格数据转为文本格式"""
        if not table_data:
            return ""

        lines = []
        for row in table_data:
            # 清理单元格内容
            cleaned_row = [str(cell).strip() if cell else "" for cell in row]
            lines.append(" | ".join(cleaned_row))

        return "\n".join(lines)

    def parse_word(self, file_path: str) -> Dict[str, str]:
        """解析 Word 文件"""
        try:
            doc = Document(file_path)
            content = []

            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)

            # 提取表格
            if doc.tables:
                content.append("\n--- 表格内容 ---")
                for table_idx, table in enumerate(doc.tables):
                    content.append(f"\n表格 {table_idx + 1}:")
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        content.append(" | ".join(row_data))

            return {
                'content': '\n'.join(content),
                'metadata': {
                    'type': 'Word',
                    'paragraphs': len(doc.paragraphs),
                    'tables': len(doc.tables),
                    'file_name': os.path.basename(file_path)
                }
            }
        except Exception as e:
            return {
                'content': '',
                'metadata': {'type': 'Word', 'error': str(e)}
            }

    def parse_excel(self, file_path: str) -> Dict[str, str]:
        """解析 Excel 文件（支持多sheet、中英文混合）"""
        try:
            # 使用 pandas 读取，支持多个工作表
            # engine='openpyxl' 支持 .xlsx, engine='xlrd' 支持 .xls
            file_ext = os.path.splitext(file_path)[1].lower()
            engine = 'openpyxl' if file_ext == '.xlsx' else 'xlrd'

            excel_file = pd.ExcelFile(file_path, engine=engine)
            content = []
            total_rows = 0

            for sheet_name in excel_file.sheet_names:
                # 读取时保持字符串格式，避免数字转换问题
                df = pd.read_excel(
                    excel_file,
                    sheet_name=sheet_name,
                    dtype=str,  # 保持原始字符串格式
                    na_filter=False  # 不将空值转为 NaN
                )

                content.append(f"\n{'='*50}")
                content.append(f"工作表: {sheet_name}")
                content.append(f"{'='*50}")

                # 转换为文本格式
                if not df.empty:
                    # 使用 to_string 保持表格格式，支持中英文对齐
                    content.append(df.to_string(index=False, max_colwidth=100))
                    total_rows += len(df)
                else:
                    content.append("(空表)")

            return {
                'content': '\n'.join(content),
                'metadata': {
                    'type': 'Excel',
                    'sheets': len(excel_file.sheet_names),
                    'sheet_names': excel_file.sheet_names,
                    'total_rows': total_rows,
                    'file_name': os.path.basename(file_path)
                }
            }
        except Exception as e:
            return {
                'content': '',
                'metadata': {'type': 'Excel', 'error': str(e)}
            }


def extract_text_from_file(file_path: str) -> str:
    """
    便捷函数：从文件提取文本

    Args:
        file_path: 文件路径

    Returns:
        提取的文本内容
    """
    parser = DocumentParser()
    result = parser.parse(file_path)
    return result.get('content', '')
